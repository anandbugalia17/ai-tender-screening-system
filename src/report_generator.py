"""
report_generator.py
--------------------
Phase 4: take the consolidated list of tender records and produce the
two business-facing outputs, laid out to match the reference tracker
format (S No, Ref/Tender ID, Country, Organization, Tender Title,
Category, Product/Service, Tender Value, Deadline, Eligibility,
Source Website, Remarks) plus the extra detail fields this pipeline
also extracts (EMD, Documents Required, Scope Summary, Status).
"""

from pathlib import Path
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

EXCEL_COLUMNS = [
    "s_no",
    "tender_id",
    "country",
    "organization",
    "tender_title",
    "category",
    "product_service",
    "tender_value",
    "emd_value",
    "submission_date",
    "eligibility_criteria",
    "documents_required",
    "source_website",
    "scope_summary",
    "remarks",
    "source_file",
    "status",
]

COLUMN_HEADERS = {
    "s_no": "S No",
    "tender_id": "Ref. No / Tender ID",
    "country": "Country",
    "organization": "Organization",
    "tender_title": "Tender Title",
    "category": "Category",
    "product_service": "Product/Service",
    "tender_value": "Tender Value",
    "emd_value": "EMD Value",
    "submission_date": "Deadline",
    "eligibility_criteria": "Eligibility",
    "documents_required": "Documents Required",
    "source_website": "Source Website",
    "scope_summary": "Scope Summary",
    "remarks": "Remarks",
    "source_file": "Source File",
    "status": "Processing Status",
}


def build_excel(records, output_path):
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    df = pd.DataFrame(records, columns=EXCEL_COLUMNS)
    df = df.rename(columns=COLUMN_HEADERS)

    df.to_excel(output_path, index=False, sheet_name="Tenders")
    return output_path


def build_docx(records, output_path):
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    title = doc.add_heading("Tender Summary Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Generated " + datetime.now().strftime("%d %b %Y, %H:%M"))
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("Total tenders processed: " + str(len(records)))

    def field(record, key, label, bold_label=True):
        p = doc.add_paragraph()
        r = p.add_run(label + ": ")
        r.bold = bold_label
        p.add_run(str(record.get(key) or "Not found"))

    for record in records:
        doc.add_heading(
            str(record.get("s_no", "")) + ". " + str(record.get("tender_title") or record.get("tender_id") or "Tender"),
            level=1,
        )

        field(record, "source_file", "Source file")
        field(record, "tender_id", "Ref. No / Tender ID")
        field(record, "country", "Country")
        field(record, "organization", "Organization")
        field(record, "category", "Category")
        field(record, "product_service", "Product/Service")
        field(record, "tender_value", "Tender Value")
        field(record, "emd_value", "EMD Value")
        field(record, "submission_date", "Deadline")
        field(record, "eligibility_criteria", "Eligibility")
        field(record, "documents_required", "Documents Required")
        field(record, "source_website", "Source Website")

        p = doc.add_paragraph()
        p.add_run("Scope Summary: ").bold = True
        doc.add_paragraph(record.get("scope_summary") or "Not available", style="List Bullet")

        p = doc.add_paragraph()
        p.add_run("Remarks: ").bold = True
        doc.add_paragraph(record.get("remarks") or "Not available", style="List Bullet")

        status = record.get("status", "")
        if status and status != "ok":
            p = doc.add_paragraph()
            run = p.add_run("Status: " + status)
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        doc.add_paragraph("-" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    return output_path